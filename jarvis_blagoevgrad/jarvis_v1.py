# vision ver 1.1 for live demo Blagoevgrad - 2025

import io
import re
import math
import pygame
import random
import spotipy
import requests
from deep_translator import GoogleTranslator
from langchain_ollama import OllamaLLM
import win32com.client as win32
from datetime import datetime, timedelta
import dateparser
from pycaw.pycaw import AudioUtilities
from pycaw.pycaw import IAudioEndpointVolume
from docx import Document
import subprocess

from jarvis_functions.gemini_vision_method import *
from jarvis_functions.call_phone_method import *
from jarvis_functions.whatsapp_messaging_method import *
from jarvis_functions.ocr_model_method import *
from jarvis_functions.shazam_method import *
from api_keys.api_keys import ELEVEN_LABS_API, GEMINI_KEY, SPOTIFY_CLIENT_ID, SPOTIFY_CLIENT_SECRET

# Initialize Pygame
pygame.init()
pygame.mixer.init()
client = ElevenLabs(api_key=ELEVEN_LABS_API)
r = sr.Recognizer()

#tv lights
WLED_IP = "192.168.10.211"

# Seting up spotify
client_id = SPOTIFY_CLIENT_ID
client_secret = SPOTIFY_CLIENT_SECRET
sp = spotipy.Spotify(auth_manager=spotipy.SpotifyOAuth(
    client_id=client_id,
    client_secret=client_secret,
    redirect_uri='http://localhost:8888/callback',
    scope='user-library-read user-read-playback-state user-modify-playback-state'))  # Scope for currently playing song

jazz_playlist_url = "spotify:playlist/60joMYdXRjtwwfyERiGu4c?si=42cc553fb755446d"

# Setting up Gemini
os.environ["GEMINI_API_KEY"] = GEMINI_KEY

genai.configure(api_key=os.environ["GEMINI_API_KEY"])

model = genai.GenerativeModel(model_name="gemini-1.5-flash")

system_instruction = (
    "Вие сте Джарвис, полезен и информативен AI асистент."
    "Винаги отговаряйте професионално и кратко, но също се дръж приятелски."
    "Поддържайте отговорите кратки, но информативни."
    "Осигурете, че всички отговори са фактологически точни и лесни за разбиране."
    "При представяне на информацията, да се има на предвид и да се адаптира за дете или тинейджър със сериозни зрителни проблеми"
)

chat = model.start_chat(
    history=[
        {
            "role": "user",
            "parts": [system_instruction],
        }
    ]
)

# Screen Dimensions
# info = pygame.display.Info()
# WIDTH, HEIGHT = info.current_w, info.current_h
# screen = pygame.display.set_mode((WIDTH, HEIGHT), pygame.FULLSCREEN)

WIDTH, HEIGHT = 1920, 1080
screen = pygame.display.set_mode((WIDTH, HEIGHT), pygame.RESIZABLE)
pygame.display.set_caption("Jarvis Interface")

# Colors
BLACK = (0, 0, 0)
WHITE = (255, 255, 255)
BLUE = (0, 128, 255)
CYAN = (0, 255, 255)
ORANGE1 = (255, 165, 0)
ORANGE2 = (255, 115, 0)
GREEN1 = (0, 219, 0)
GREEN2 = (4, 201, 4)
PINK1 = (255, 182, 193)  # Light Pink
PINK2 = (255, 105, 180)  # Hot Pink
PURPLE1 = (166, 0, 255)
PURPLE2 = (176, 28, 255)
font_large = pygame.font.Font(None, 48)
font_small = pygame.font.Font(None, 32)

# Fonts
font_large = pygame.font.Font(pygame.font.get_default_font(), 36)
font_small = pygame.font.Font(pygame.font.get_default_font(), 20)

# Clock
clock = pygame.time.Clock()

# Rotating Circle Parameters
center = (WIDTH // 2, HEIGHT // 2)
max_radius = min(WIDTH, HEIGHT) // 3
angle = 0
speed = 1

# Particle Parameters
particles = []
num_particles = 100

# Pulse effect variables
pulse_factor = 1
pulse_speed = 0.05
min_size = 3
max_size = 3

# Color Transition
current_color_1 = list(BLUE)
current_color_2 = list(CYAN)
target_color_1 = list(BLUE)
target_color_2 = list(CYAN)
color_transition_speed = 10

jarvis_responses = [
    "Тук съм, как мога да помогна?",
    "Слушам, как мога да Ви асистирам?",
    "Тук съм, как мога да помогна?",
    "С какво мога да Ви бъда полезен?",
    "Слушам шефе, как да помогна?"
]

selected_songs = [
    "Another one bites the dust - Queen",
    "Back in black",
    "Shoot to Thrill",
    "Thunderstruck",
    "You Give Love a Bad Name",
    "Highway to Hell - AC/DC",
    "September - Earth, Wind & Fire",
    "Should I Stay or Should I Go - Remastered",
    "If You Want Blood(You've Got It) - AC/DC",
    "Welcome Tо The Jungle - Guns N' Roses"
]

status_list = []

jarvis_voice = "Brian" #deffault voice

# Ball initial random positions
random_particles = [{"x": random.randint(0, WIDTH), "y": random.randint(0, HEIGHT),
                     "dx": random.uniform(-2, 2), "dy": random.uniform(-2, 2)} for _ in range(num_particles)]

# State Variables
model_answering = False
is_collided = False
is_generating = False
wake_word_detected = False

running = True
current_song = ""
current_artist = ""
album_cover = None
current_progress = 0
song_duration = 0

def get_current_volume():
    # Get the default audio device
    devices = AudioUtilities.GetSpeakers()
    interface = devices.Activate(
        IAudioEndpointVolume._iid_, 0, None)
    volume = interface.QueryInterface(IAudioEndpointVolume)

    # Get the current volume level (0.0 to 1.0)
    current_volume = volume.GetMasterVolumeLevelScalar()
    return current_volume * 100  # Return as percentage

def set_volume(percentage):
    # Convert percentage to a value between 0.0 and 1.0
    if 0 <= percentage <= 100:
        volume = AudioUtilities.GetSpeakers().Activate(
            IAudioEndpointVolume._iid_, 0, None).QueryInterface(IAudioEndpointVolume)
        volume.SetMasterVolumeLevelScalar(percentage / 100.0, None)
        print(f"Volume set to {percentage}%")
    else:
        print("Please provide a valid volume percentage between 0 and 100.")

def mute(unmute=False):
    # Mute or unmute the audio
    volume = AudioUtilities.GetSpeakers().Activate(
        IAudioEndpointVolume._iid_, 0, None).QueryInterface(IAudioEndpointVolume)
    volume.SetMute(unmute, None)
    print("Audio is muted." if unmute else "Audio is unmuted.")

def send_email(subject, body, to_email):
    outlook = win32.Dispatch('outlook.application')
    mail = outlook.CreateItem(0)
    mail.Subject = subject
    mail.Body = body
    mail.To = to_email
    mail.Send()

def parse_natural_time(natural_time):
    """
    Parses a natural language time expression (e.g., '3 часа следобяд днес', 'tomorrow', 'next Wednesday')
    into a datetime object.
    """

    # Manually handle 'днес' and 'утре' since dateparser fails sometimes
    now = datetime.now()

    # Replace Bulgarian words with English for better parsing
    normalized_time = (
        natural_time.replace("днес", "today")
        .replace("утре", "tomorrow")
        .replace("следобяд", "PM")
        .replace("сутринта", "AM")
    )

    # Try parsing with dateparser
    event_time = dateparser.parse(
        normalized_time,
        languages=['bg', 'en'],  # Use both Bulgarian and English
        settings={'PREFER_DATES_FROM': 'future'}
    )

    # If dateparser fails, manually handle simple cases
    if event_time is None:
        if "днес" in natural_time:
            event_time = now.replace(hour=15, minute=0, second=0, microsecond=0)
        elif "утре" in natural_time:
            event_time = (now + timedelta(days=1)).replace(hour=15, minute=0, second=0, microsecond=0)
        else:
            raise ValueError(f"Could not parse the given time expression: {natural_time}")

    return event_time

def create_outlook_appointment(subject, start_time, duration):
    outlook = win32.Dispatch("Outlook.Application")
    appointment = outlook.CreateItem(1)  # 1 = olAppointmentItem

    appointment.Subject = subject
    appointment.Start = start_time
    appointment.Duration = duration
    appointment.ReminderMinutesBeforeStart = 15
    appointment.Save()

    print(f"✅ Appointment '{subject}' scheduled for {start_time}")

def blend_color(current, target, speed):
    """Gradually transitions the current color toward the target color."""
    for i in range(3):
        diff = target[i] - current[i]
        if abs(diff) > speed:
            current[i] += speed if diff > 0 else -speed
        else:
            current[i] = target[i]

def draw_particles(surface, particles, target_mode=False):
    """Draws particles on the surface. If target_mode is True, arrange them in a circle and pulse."""
    global angle, pulse_factor

    for i, particle in enumerate(particles):
        if target_mode:
            # Calculate target circular positions
            target_x = center[0] + math.cos(math.radians(angle + i * 360 / len(particles))) * max_radius
            target_y = center[1] + math.sin(math.radians(angle + i * 360 / len(particles))) * max_radius
            # Smoothly move particles towards their circular positions
            particle["x"] += (target_x - particle["x"]) * 0.05
            particle["y"] += (target_y - particle["y"]) * 0.05

            # Pulse effect
            pulse_factor = min(max_size, pulse_factor + pulse_speed) if pulse_factor < max_size else max(min_size, pulse_factor - pulse_speed)
        else:
            # Move particles randomly when in default mode
            particle["x"] += particle["dx"]
            particle["y"] += particle["dy"]

            # Keep particles within the screen bounds
            if particle["x"] <= 0 or particle["x"] >= WIDTH:
                particle["dx"] *= -1
            if particle["y"] <= 0 or particle["y"] >= HEIGHT:
                particle["dy"] *= -1

        # Draw the particle
        pygame.draw.circle(surface, tuple(current_color_2), (int(particle["x"]), int(particle["y"])), int(pulse_factor))

def draw_response(model):
    """Update settings when the model is answering."""
    global target_color_1, target_color_2, is_collided, angle, speed

    if model == "Jarvis":
        target_color_1 = list(GREEN1)
        target_color_2 = list(GREEN2)
    elif model == "Friday":
        target_color_1 = list(PINK1)
        target_color_2 = list(PINK2)
    elif model == "Veronica":
        target_color_1 = list(PURPLE1)
        target_color_2 = list(PURPLE2)

    speed = 1
    is_collided = True
    angle += speed

def draw_thinking():
    """Update settings when the model is listening."""
    global target_color_1, target_color_2, is_collided, angle, speed
    target_color_1 = list(ORANGE1)
    target_color_2 = list(ORANGE2)
    speed = 0.5
    is_collided = True
    angle += speed

def draw_default():
    """Update settings when the model is not answering."""
    global target_color_1, target_color_2, is_collided, speed
    target_color_1 = list(BLUE)
    target_color_2 = list(CYAN)
    speed = 1
    is_collided = False

def draw_text(surface, text, position, font, color):
    """Draws text onto the surface."""
    text_surface = font.render(text, True, color)
    surface.blit(text_surface, position)

def extract_keywords(user_input):
    """
    Extract keywords from a natural language query.
    In this case, remove "can you search for" and return the rest.
    """
    trigger_phrase = "Може ли да потърсиш за"
    if user_input.lower().startswith(trigger_phrase):
        # Extract the part after the trigger phrase
        return user_input[len(trigger_phrase):].strip()
    return None

def perform_web_search(search_term):
    """
    Perform a web search using the extracted keywords.
    """
    if not search_term:
        print("No valid search term provided.")
        return
    # Encode the search term for the URL
    encoded_term = search_term.replace(" ", "+")
    # Construct the PowerShell command
    command = f'Start-Process "firefox.exe" "https://www.google.com/search?q={encoded_term}"'
    # Execute the PowerShell command
    subprocess.run(["powershell", "-Command", command], shell=True)

def fetch_current_track():
    """Fetch the current playing track and its album cover."""
    try:
        current_track = sp.currently_playing()
        if current_track and current_track['is_playing']:
            song = current_track['item']['name']
            artist = ", ".join([a['name'] for a in current_track['item']['artists']])
            album_cover_url = current_track['item']['album']['images'][0]['url']
            progress_ms = current_track['progress_ms']  # Progress in milliseconds
            duration_ms = current_track['item']['duration_ms']  # Duration in milliseconds
            return song, artist, album_cover_url, progress_ms, duration_ms
        return None, None, None, 0, 0
    except Exception as e:
        print(f"Error fetching track: {e}")
        return None, None, None, 0, 0

def load_album_cover(url):
    """Download and convert the album cover image to a Pygame surface."""
    try:
        response = requests.get(url)
        if response.status_code == 200:
            image_data = io.BytesIO(response.content)
            image = pygame.image.load(image_data, "jpg")
            return pygame.transform.scale(image, (300, 300))  # Scale to 300x300
    except Exception as e:
        print(f"Error loading album cover: {e}")
    return None

def play_music():
    sp.start_playback()  # Start playback (Play the song)

def pause_music():
    sp.pause_playback()  # Pause the playback (Stop the song)

def draw_progress_bar(surface, x, y, width, height, progress, max_progress):
    """Draw a progress bar to represent the song timeline."""
    # Check if max_progress is non-zero to avoid division by zero
    if max_progress > 0:
        # Calculate the progress ratio
        progress_ratio = progress / max_progress
        progress_width = int(width * progress_ratio)
    else:
        progress_width = 0  # If duration is zero, show no progress

    # Draw the empty progress bar (background)
    pygame.draw.rect(surface, (50, 50, 50), (x, y, width, height))

    # Draw the filled progress bar (foreground)
    pygame.draw.rect(surface, GREEN1, (x, y, progress_width, height))

def set_color(red, green, blue):
    """Set the color using RGB values."""
    url = f"http://{WLED_IP}/json/state"
    data = {
        "on": True,
        "bri": 255,  # Optional: brightness
        "seg": [{
            "col": [[red, green, blue]]
        }]
    }
    response = requests.post(url, json=data)

def update_status(new_status):
    # Add new status to the list
    status_list.append(new_status)

    # Ensure the list only has 5 elements
    if len(status_list) > 5:
        status_list.pop(0)  # Remove the oldest status (first element)

def record_text():
    """Listen for speech and return the recognized text."""
    try:
        with sr.Microphone() as source:
            #print("Listening...")
            r.adjust_for_ambient_noise(source, duration=0.2)
            audio = r.listen(source)

            # Recognize speech using Google API
            MyText = r.recognize_google(audio, language="bg-BG")
            print(f"You said: {MyText}")
            return MyText.lower()

    except sr.RequestError as e:
        print(f"API Request Error: {e}")
        return None
    except sr.UnknownValueError:
        print("Sorry, I didn't catch that. Please try again.")
        return None

def chatbot():
    global wake_word_detected, model_answering, is_generating, current_model

    current_model= "Jarvis"

    print("Welcome to Vision! Say any of the models name to activate. Say 'exit' to quit.")

    while True:
        if not wake_word_detected:
            # Listen for the wake word
            print("Waiting for wake word...")
            user_input = record_text()

            if user_input and ("джарвис" in user_input or "джарви" in user_input or "джервис" in user_input):
                wake_word_detected = True
                current_model = "Jarvis"
                pygame.mixer.music.load("sound_files/beep.flac")
                pygame.mixer.music.play()

                print("✅ Wake word detected!")
                model_answering = True
                is_generating = False

                jarvis_voice = "Brian"
                response = random.choice(jarvis_responses)
                audio = client.generate(text=response, voice=jarvis_voice)
                play(audio)

                model_answering = False
                is_generating = True

            elif user_input and "friday" in user_input:
                wake_word_detected = True
                current_model = "Friday"
                pygame.mixer.music.load("sound_files/beep.flac")
                pygame.mixer.music.play()

                print("✅ Wake word detected!")
                model_answering = True
                is_generating = False

                jarvis_voice = "Matilda"
                response = random.choice(jarvis_responses)
                audio = client.generate(text=response, voice=jarvis_voice)
                play(audio)

                model_answering = False
                is_generating = True

            elif user_input and ("Вероника" in user_input or "вероника" in user_input or "Veronica" in user_input):
                wake_word_detected = True
                current_model = "Veronica"
                pygame.mixer.music.load("sound_files/beep.flac")
                pygame.mixer.music.play()

                print("✅ Wake word detected!")
                model_answering = True
                is_generating = False

                jarvis_voice = "Sarah"
                response = random.choice(jarvis_responses)
                audio = client.generate(text=response, voice=jarvis_voice)
                play(audio)

                model_answering = False
                is_generating = True

            elif user_input == "излез":
                print("Goodbye!")
                jarvis_voice = "Brian"
                audio = client.generate(text="Goodbye!", voice=jarvis_voice)
                play(audio)
                break

        else:
            # Actively listen for commands
            print("Listening for commands...")
            user_input = record_text()
            if user_input is None:
                print("Error: No input detected.")
                continue

            if "целта" in user_input and "проекта" in user_input:
                audio = client.generate(
                    text="Проектът VISION е създаден с мисията да разработи интелигентен гласов асистент, който да подпомага деца със зрителни и други затруднения в работата им с компютри. "
                         "Основната му роля е да улесни образователния процес, личностното развитие и социалната интеграция на тези деца."
                          "Чрез внедряване на технологии за разпознаване на реч и текст на български език, системата ще осигури по-лесен достъп до информация, ще подпомогне комуникацията и ще позволи управление на умни устройства, като по този начин ще допринесе за по-добро качество на живот.",
                    voice="Brian")
                play(audio)
                model_answering = False
                is_generating = False
                wake_word_detected = False
                continue

            if "кои" in user_input and "потребителите" in user_input:
                audio = client.generate(
                    text="Въпреки че VISION е създаден с фокус върху деца със зрителни и други затруднения, той е полезен за всеки, който търси интелигентен гласов асистент с усъвършенствани възможности за комуникация, автоматизация на офис задачи и интеграция с големи езикови модели."
                        "Основните групи потребители включват:"
                        "Хора със зрителни и други увреждания, Ученици и студенти, Възрастни хора и Всички, които желаят удобен и функционален гласов асистент в ежедневието си",
                    voice="Brian")
                play(audio)
                model_answering = False
                is_generating = False
                wake_word_detected = False
                continue

            if ("кои" in user_input or "Кои" in user_input) and "инструменти" in user_input:
                audio = client.generate(
                    text="Иструментите използвани във Vision са следните: За програмиране на програмата е използван Python, Pycharm и PyGame."
                         "Като също е използван OpenCV, EasyOCR и Gemini Vision за разпознаване на изображения"
                         "За разпознаването на гласа е използвана библиотеката Speech Recognition, а за възпроизвеждането на гласа който чувате в момента е използван Eleven Labs."
                         "Отделно са използвани библиотеки които могат да управляват Whatsapp, Word и Spotify"
                        "А моделът на който е базиран Jarvis е Gemini едно цяло и пет Flash, като във бъдеще ще има опция за DeepSeek и Llama 3",
                    voice="Brian")
                play(audio)
                model_answering = False
                is_generating = False
                wake_word_detected = False
                continue


            if "представи се" in user_input or "представиш" in user_input:
                audio = client.generate(text="Здравейте, аз съм Джарвис, акроним от (Just A Rather Very Intelligent System), аз съм езиков модел на Gemini обучен от Google."
                                             "Вдъхновен съм от легендарния изкуствен интелект на Тони Старк – Джарвис от Железния човек."
                                              "Аз съм тук, за да отговоря на въпросите ви, да помогна със задачи или да водя разговори на всякакви теми. "
                                             "Ако искате да ме попитате нещо, просто ме повикайте по име.", voice="Brian")
                play(audio)
                model_answering = False
                is_generating = False
                continue

            if "можеш" in user_input and "правиш" in user_input:
                audio = client.generate(text="Мога да търся информация в интернет, да я обобщавам и да ви я представям. "
                                             "Също така, мога да изпращам и чета имейли, да пускам музика, да отварям нови документи в Word "
                                             "И дори да ви опиша това, което виждам като изпозлвам Gemini Vision и OCR модел за разпознаване на текст.",
                                        voice="Brian")
                play(audio)
                model_answering = False
                is_generating = False
                continue


            if ("пусни" in user_input or "пуснеш" in user_input) and ("песен" in user_input or "музика" in user_input):
                model_answering = True
                is_generating = False

                audio = client.generate(text="Разбира се, имате ли някакви предпочитания за песен?", voice=jarvis_voice)
                play(audio)

                model_answering = False
                is_generating = True

                print("Listening for song info...")
                user_input = record_text()

                if user_input is None:
                    audio = client.generate(text="Нo можах да разбера. Може ли да повторите?", voice=jarvis_voice)
                    play(audio)
                    user_input = record_text()

                if "да" in user_input:
                    model_answering = True
                    is_generating = False

                    audio = client.generate(text="Добре, коя песен бихте желали да ви пусна?",
                                            voice=jarvis_voice)
                    play(audio)

                    print("Listening for specific song...")
                    user_input = record_text()

                    audio = client.generate(text=f"Пускам, {user_input}",
                                            voice=jarvis_voice)
                    play(audio)

                    track_name = user_input
                    result = sp.search(q=track_name, limit=1)

                    # Get the song's URI
                    track_uri = result['tracks']['items'][0]['uri']
                    print(f"Playing track: {track_name}")

                    # Get the current device
                    devices = sp.devices()
                    # Find the LAPTOP_KOSI device by its ID
                    pc_device_id = '7993e31456b6d73672f9c7bcee055fb10ae52f23'
                    update_status(f"Played {track_name}")

                    # Start playback on the LAPTOP_KOSI device
                    sp.start_playback(device_id=pc_device_id, uris=[track_uri])
                    print("Playback started on LAPTOP_KOSI.")

                elif "не" in user_input:
                    model_answering = True
                    is_generating = False

                    audio = client.generate(text="Пускам тогава от избрания от вас списък с песни?",
                                            voice=jarvis_voice)
                    play(audio)

                    track_name = random.choice(selected_songs)
                    result = sp.search(q=track_name, limit=1)

                    # Get the song's URI
                    track_uri = result['tracks']['items'][0]['uri']
                    print(f"Playing track: {track_name}")

                    # Get the current device
                    devices = sp.devices()
                    # Find the LAPTOP_KOSI device by its ID
                    pc_device_id = '7993e31456b6d73672f9c7bcee055fb10ae52f23'
                    update_status(f"Played {track_name}")

                    # Start playback on the LAPTOP_KOSI device
                    sp.start_playback(device_id=pc_device_id, uris=[track_uri])
                    print("Playback started on LAPTOP_KOSI.")

                model_answering = False
                is_generating = False
                wake_word_detected = False
                continue

            if "спри" in user_input and ("песента" in user_input or "музиката" in user_input):
                pause_music()
                model_answering = False
                is_generating = False
                wake_word_detected = False
                continue

            if "пратиш" in user_input and ("имейл" in user_input or "писмо" in user_input):
                model_answering = True
                is_generating = False

                audio = client.generate(text="Разбира се, към кого бихте желали да пратите имейла?", voice=jarvis_voice)
                play(audio)

                print("Listening for email info...")
                user_input = record_text()

                if "тати" in user_input or "баща ми" in user_input:
                    to_email = "bojidarbojinov@outlook.com"
                elif "мама" in user_input or "майка ми" in user_input:
                    to_email = "kameliqbojinova@outlook.com"

                audio = client.generate(text="Каква ще е темата на вашето писмо?", voice=jarvis_voice)
                play(audio)

                print("Listening for email info...")
                subject = record_text()

                audio = client.generate(text="Какво искате да изпратите?", voice=jarvis_voice)
                play(audio)

                print("Listening for email info...")
                body = record_text()

                audio = client.generate(text="Супер, преди да изпратя имейла, ще ви кажа какво съм си записал",
                                        voice=jarvis_voice)
                play(audio)

                if to_email == "bojidarbojinov@outlook.com":
                    audio = client.generate(text="Имейла е към Божидар Божинов (баща ви)", voice=jarvis_voice)
                    play(audio)
                elif to_email == "kameliqbojinova@outlook.com":
                    audio = client.generate(text="Имейла е към Камелия Божинова (майка ви)", voice=jarvis_voice)
                    play(audio)
                audio = client.generate(text="Темата на писмото е " + subject + "И съдържанието на писмото е " + body,
                                        voice=jarvis_voice)
                play(audio)

                audio = client.generate(text="Всичко наред ли е с информацията в писмото?", voice=jarvis_voice)
                play(audio)

                print("Listening for approval...")
                user_input = record_text()

                if "да" in user_input:
                    audio = client.generate(text="✅ Супер, пращам имейла", voice=jarvis_voice)
                    play(audio)
                    send_email(subject, body, to_email)
                    update_status(f"Sent an email to {to_email}")

                elif "не" in user_input:
                    audio = client.generate(text="Сорка", voice=jarvis_voice)
                    play(audio)

                model_answering = False
                is_generating = False
                wake_word_detected = False
                continue

            if "прочетеш" in user_input and ("писма" in user_input or "имейли" in user_input or "пис" in user_input):
                # Initialize Outlook
                outlook = win32.Dispatch("Outlook.Application").GetNamespace("MAPI")
                inbox = outlook.GetDefaultFolder(6)  # 6 = Inbox

                # Get all messages sorted by received time (newest first)
                messages = inbox.Items
                messages.Sort("[ReceivedTime]", True)  # Sort descending (newest first)

                # Retrieve the last 5 emails
                num_emails = 3  # Change this number if you need more
                latest_messages = [messages.GetNext() for _ in range(num_emails)]

                audio = client.generate(text="Ето последните 3 имейла в пощата ви: ", voice=jarvis_voice)
                play(audio)
                # Print email details
                for i, email in enumerate(latest_messages, start=1):
                    print(f"\n📧 Email {i}:")
                    print(f"Subject: {email.Subject}")
                    print(f"From: {email.SenderName}")
                    print(f"Received: {email.ReceivedTime}")
                    print("\n--- Email Body ---\n")
                    print(email.Body)  # Full email body
                    print("\n--- End of Email ---\n")
                    audio = client.generate(text=f"Имейл номер {i}, изпратено е от {email.SenderName}, "
                                                 f"темата е {email.Subject}, а съдържанието на писмото е {email.Body}", voice=jarvis_voice)
                    play(audio)

                update_status(f"Read last 3 emails")
                model_answering = False
                is_generating = False
                wake_word_detected = False
                continue

            if (("събитие" in user_input or "събити" in user_input or "събития" in user_input)
                    and ("създадеш" in user_input or "Създадеш" in user_input or "създай" in user_input or "Създай" in user_input)):
                # subject of event
                audio = client.generate(text="Разбира се, как искате да се казва събитието?", voice=jarvis_voice)
                play(audio)

                print("Listening for apointment info...")
                subject = record_text()

                # time of event
                audio = client.generate(text="За кога да бъде това събитие?", voice=jarvis_voice)
                play(audio)

                print("Listening for apointment info...")
                user_input = record_text()

                # duration of event
                audio = client.generate(text="Колко време ще продължи това събитие?", voice=jarvis_voice)
                play(audio)

                print("Listening for apointment info...")
                duration = record_text()

                try:
                    event_time = parse_natural_time(user_input)
                    print(f"Parsed event time: {event_time}")  # Debug output
                    audio = client.generate(
                        text=f"Супер, запазвам събитие {subject}, в {event_time.strftime('%H:%M %d-%m-%Y')}, и ще трае 1 час",
                        voice=jarvis_voice)
                    play(audio)
                    create_outlook_appointment(subject, event_time, duration = 60)
                    update_status(f"Made an event in the calendar")
                    model_answering = False
                    is_generating = False
                    wake_word_detected = False
                    continue
                except ValueError as e:
                    print(f"❌ Error: {e}")

                # Направи ми събитие за 3 следобяд днес, което да продължи 1 час, и да се казва "нахрани котката"pip install pywin32

            if ("виждаш" in user_input or "вижда" in user_input) and "какво" in user_input:
                gemini_vision()
                update_status(f"Used Gemini Vision")
                model_answering = False
                is_generating = False
                wake_word_detected = False
                continue

            if ("пише" in user_input or "пиша" in user_input) and "какво" in user_input:
                # Open the webcam
                audio = client.generate(text="Камерата по подразбиране ли да използвам?", voice=jarvis_voice)
                play(audio)

                print("Listening for camera info...")
                camera_info = record_text()

                camera_index = 0 #deffault index

                if "да" in camera_info:
                    camera_index = 0
                    audio = client.generate(text="Добре, използвам web камерата на компютъра ви", voice=jarvis_voice)
                    play(audio)
                elif "не" in camera_info or "другата" in user_input:
                    camera_index = 1
                    audio = client.generate(text="Добре, използвам камерата от ви ар хедсета",
                                            voice=jarvis_voice)
                    play(audio)

                extracted_text_from_ocr = capture_and_ocr(camera_index = camera_index)

                if len(extracted_text_from_ocr) > 0:
                    print("Here is what the OCR model detected:")
                    print(extracted_text_from_ocr)
                    audio = client.generate(text=fr"OCR модела разпозна следното {extracted_text_from_ocr}",
                                            voice=jarvis_voice)
                    play(audio)
                else:
                    print("Nothing was detected by the OCR model")
                    audio = client.generate(text="OCR модела не можа да разпознае нищо от снимката",
                                            voice=jarvis_voice)
                    play(audio)

                update_status(f"Used the OCR model")
                model_answering = False
                is_generating = False
                wake_word_detected = False
                continue

            if "звъннеш" in user_input:
                call_phone()

                update_status(f"Called Tati")
                model_answering = False
                is_generating = False
                wake_word_detected = False
                continue

            if ("съобщение" in user_input or "съобщения" in user_input) and "пратиш" in user_input:
                whatsapp_send_message()

                update_status(f"Sent message to Tati")
                model_answering = False
                is_generating = False
                wake_word_detected = False
                continue

            if ("разпознаеш" in user_input or "коя" in user_input) and "песен" in user_input:
                audio = client.generate(text="Разбира се, започвам да слушам. Ако разпозная песента ще ви кажа името и автора на песента",
                                        voice=jarvis_voice)
                play(audio)

                title, artist = recognize_audio()  # Get the title and artist
                if title and artist:
                    audio = client.generate(
                        text=f"Намерих песента, песента е {title}, а автора е {artist}. Желаете ли да пусна песента в spotify?",
                        voice=jarvis_voice)
                    play(audio)
                    print(f"Song Title: {title}")
                    print(f"Artist: {artist}")

                    print("Listening for song info...")
                    answer_info = record_text()

                    if "да" in answer_info:
                        audio = client.generate(text=f"Пускам, {title} на {artist}",
                                                voice=jarvis_voice)
                        play(audio)
                        track_name = {title}
                        result = sp.search(q=track_name, limit=1)

                        # Get the song's URI
                        track_uri = result['tracks']['items'][0]['uri']
                        print(f"Playing track: {track_name}")

                        # Get the current device
                        devices = sp.devices()
                        # Find the LAPTOP_KOSI device by its ID
                        pc_device_id = '7993e31456b6d73672f9c7bcee055fb10ae52f23'
                        update_status(f"Played {track_name}")

                        # Start playback on the LAPTOP_KOSI device
                        sp.start_playback(device_id=pc_device_id, uris=[track_uri])
                        print("Playback started on LAPTOP_KOSI.")

                    elif "не" in answer_info:
                        model_answering = False
                        is_generating = False
                        wake_word_detected = False
                        continue
                else:
                    print("No song found")

                update_status(f"Used Shazam")
                model_answering = False
                is_generating = False
                wake_word_detected = False
                continue

            if (("отвори" in user_input or "отвориш" in user_input or "отвориш" in user_input ) # currently not working
                    and ("word" in user_input or "wor" in user_input or "документ" in user_input)):
                # Initialize document
                doc = Document()

                # Open Word
                audio = client.generate(text="Разбира се, отварям Word. Само секунда", voice=jarvis_voice)
                play(audio)

                # Wait a bit to ensure Word has opened
                time.sleep(2)

                # Ask for document title
                audio = client.generate(
                    text="Готов съм. Преди да започнем, как ще желаете да е заглавието на документа?",
                    voice=jarvis_voice)
                play(audio)

                # Listen for the title input
                print("Listening for title...")
                input_text = record_text()

                # Add a title
                doc.add_heading(input_text, 0)

                # Inform the user
                audio = client.generate(
                    text="Добре започвам да слушам и записвам. Кажете думата Край за да спра да записвам",
                    voice=jarvis_voice)
                play(audio)

                words_in_document = ""

                while True:
                    with sr.Microphone() as source:
                        try:
                            print("Listening for input...")
                            input_text = record_text()
                            print(f"You said: {input_text}")

                            # Skip if input_text is None (empty or unrecognized speech)
                            if input_text is None or input_text.strip() == "":
                                print("No speech detected or input is empty, try again.")
                                continue  # Skip to the next loop iteration

                            # Stop listening when "край" is said
                            if "край" in input_text or "Край" in input_text:
                                audio = client.generate(text="Спрях да записвам, файла е запазен в папка Downloads",
                                                        voice=jarvis_voice)
                                play(audio)
                                # Add a paragraph
                                doc.add_paragraph(words_in_document)
                                break

                            # Append the valid input to the document
                            words_in_document += input_text + ". "

                            time.sleep(1)  # Slight delay for realism

                        except sr.UnknownValueError:
                            print("Could not understand, try again.")
                        except sr.RequestError:
                            print("Speech recognition service error.")

                # Finished
                print("Document saved and process ended.")

                file_path = r'D:\example.docx'
                doc.save(file_path)
                os.system(f'start {file_path}')  # Open Word on Windows

                update_status(f"Made a Word document")
                model_answering = False
                is_generating = False
                wake_word_detected = False
                continue

            if "намали" in user_input and "звука" in user_input:
                current_volume = get_current_volume()
                print(f"Current volume: {current_volume}%")

                # Set volume to 50%
                set_volume(50)

                update_status(f"Set volume to 50%")
                model_answering = False
                is_generating = False
                wake_word_detected = False
                continue

            if "усили" in user_input and "звука" in user_input:
                current_volume = get_current_volume()
                print(f"Current volume: {current_volume}%")

                # Set volume to 50%
                set_volume(75)

                update_status(f"Set volume to 75%")
                model_answering = False
                is_generating = False
                wake_word_detected = False
                continue

            if "заглуши" in user_input and "звука" in user_input:
                current_volume = get_current_volume()
                print(f"Current volume: {current_volume}%")

                # Mute the audio
                mute(unmute=True)

                update_status(f"Mutted volume")
                model_answering = False
                is_generating = False
                wake_word_detected = False
                continue

            if "отглуши" in user_input and "звука" in user_input:
                current_volume = get_current_volume()
                print(f"Current volume: {current_volume}%")

                # Mute the audio
                mute(unmute=False)

                update_status(f"Unmuted volume")
                model_answering = False
                is_generating = False
                wake_word_detected = False
                continue

            if user_input:
                # Start thinking state
                is_generating = True

                if (current_model == "Jarvis"): #Jarvis model (Gemini)
                    result = chat.send_message({"parts": [user_input]})

                elif (current_model == "Friday"):  # Friday model (Llama3)
                    model = OllamaLLM(model="llama3")

                    translated_input_bg_to_en = (GoogleTranslator(source='bg', target='en')
                                                 .translate(user_input))
                    print(translated_input_bg_to_en)
                    result = model.invoke(input=translated_input_bg_to_en)

                elif (current_model == "Veronica"): #Friday model (Llama3)
                    model = OllamaLLM(model="deepseek-r1:1.5b")

                    translated_input_bg_to_en = (GoogleTranslator(source='bg', target='en')
                                                 .translate(user_input))
                    print(translated_input_bg_to_en)

                    # Build the full input for the model with the translated text
                    full_input = f"{system_instruction}\n\nUser: {translated_input_bg_to_en}\nAssistant:"

                    # Get the model result
                    result1 = model.invoke(input=full_input)

                    # Remove the <think> part
                    result = re.sub(r"<think>.*?</think>", "", result1, flags=re.DOTALL)

                    print(result)  # For testing

                # Done generating the answer
                is_generating = False
                model_answering = True

                # Answering based on model
                if (current_model == "Jarvis"): #Jarvis answering
                    print(f"Jarvis: {result.text}")
                    audio = client.generate(text=result.text, voice=jarvis_voice)
                    play(audio)

                elif (current_model == "Friday"):  # Friday answering
                    print(f"FRIDAY: {result}")

                    translated_text_en_to_bg = GoogleTranslator(source='en', target='bg').translate(result)
                    print(translated_text_en_to_bg)

                    # Generate audio from the translated text
                    audio = client.generate(text=translated_text_en_to_bg, voice=jarvis_voice)

                    # Play the generated audio
                    play(audio)

                elif (current_model == "Veronica"):  # Friday answering
                    print(f"Veronica: {result}")

                    translated_text_en_to_bg = GoogleTranslator(source='en', target='bg').translate(result)
                    print(translated_text_en_to_bg)  # Output: "Здравей, как си?"

                    # Generate audio from the translated text
                    audio = client.generate(text=translated_text_en_to_bg, voice=jarvis_voice)

                    # Play the generated audio
                    play(audio)

                model_answering = False
                is_generating = False

            wake_word_detected = False
# Main Loop
running = True
chatbot_thread = None

# Run chatbot in a separate thread
import threading
chatbot_thread = threading.Thread(target=chatbot)
chatbot_thread.start()

while running:
    screen.fill(BLACK)

    # Event Handling
    keys = pygame.key.get_pressed()
    for event in pygame.event.get():
        if event.type == pygame.QUIT:
            running = False

    # Toggle behavior based on whether the model is generating or answering
    if is_generating:
        draw_thinking()  # Show thinking state
        #set_color(255, 165, 0)  # Orange
    elif model_answering:
        draw_response(current_model) # Show answering state
        #set_color(0, 219, 0)  # Green
    else:
        draw_default()  # Default state when nothing is happening.
        #set_color(0, 128, 255)  # Red

    # Smooth Color Transition
    blend_color(current_color_1, target_color_1, color_transition_speed)
    blend_color(current_color_2, target_color_2, color_transition_speed)

    # Draw Particles
    draw_particles(screen, random_particles, target_mode=is_collided)

    # Draw Text
    draw_text(screen, "Vision Interface", (10, 10), font_large, WHITE)
    draw_text(screen, "System Status: All Systems Online", (10, 60), font_small, tuple(current_color_2))

    # Draw the list of statuses under "System Status"
    start_y = 90  # Starting position for the list of items
    line_height = 30  # Space between each list item

    for index, status in enumerate(status_list):
        draw_text(screen, status, (10, start_y + index * line_height), font_small, WHITE)

    # Function to update the status list
    def update_status(new_status):
        # Add new status to the list and remove the oldest one if needed
        status_list.append(new_status)
        if len(status_list) > 5:  # Keep the list size manageable
            status_list.pop(0)

    # Fetch current track periodically (e.g., every 3 seconds)
    if pygame.time.get_ticks() % 3000 < 50:  # Update every 3 seconds
        song, artist, album_cover_url, progress_ms, duration_ms = fetch_current_track()
        if song and artist:  # Only update if song and artist are available
            current_song = song
            current_artist = artist
            current_progress = progress_ms
            song_duration = duration_ms

        # Draw the progress bar for the song timeline
        # Adjust the progress bar position if needed
    progress_bar_x = (WIDTH - 700) // 2
    progress_bar_y = HEIGHT - 30  # Adjust y-position for progress bar
    draw_progress_bar(screen, progress_bar_x, progress_bar_y, 700, 10, current_progress, song_duration)

    # Draw song information above the progress bar
    if current_song:
        song_surface = font_small.render(current_song, True, WHITE)
        song_text_x = (WIDTH - song_surface.get_width()) // 2
        song_text_y = progress_bar_y - 30  # Adjust y-position for song name
        screen.blit(song_surface, (song_text_x, song_text_y))

    # Update Display
    pygame.display.flip()
    clock.tick(60)

# Quit Pygame
pygame.quit()

